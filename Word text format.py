import docx
import os
# import win32com.client
from docx import Document
import comtypes.client
import comtypes.gen
from win32com import client as wc
import docx2python


docFiledicList=[]
docFiledicDocxList=[]

def screenDocFile():
# 遍历当前目录下的所有文件夹和子文件夹
    for root, dirs, files in os.walk('.'):
        # 遍历当前文件夹下的所有文件
        for file in files:
            # 检查文件扩展名是否为".doc"或".docx"
            if file.endswith('.doc'):
                # 打开Word文档
                docFiledicList.append(os.path.abspath(os.path.join(root, file)))

                # doc = docx.Document(os.path.join(root, file))
                # 在这里进行其他的操作，例如读取文本、修改文本等

def screenDocxFile():
# 遍历当前目录下的所有文件夹和子文件夹
    for root, dirs, files in os.walk('.'):
        # 遍历当前文件夹下的所有文件
        for file in files:
            # 检查文件扩展名是否为".doc"或".docx"
            if file.endswith('.docx'):
                # 打开Word文档
                docFiledicDocxList.append(os.path.abspath(os.path.join(root, file)))

                # doc = docx.Document(os.path.join(root, file))
                # 在这里进行其他的操作，例如读取文本、修改文本等


def save_doc_to_docx(rawpath):  # doc转docx
    word = wc.Dispatch("Word.Application")
    doc = word.Documents.Open(docFiledicList[0])
    output_file = os.path.splitext(rawpath)[0] + '.docx'

    doc.SaveAs(output_file, 12)  # 12表示docx格式
    doc.Close()
    word.Quit()
    word.Quit()


if __name__ == '__main__':    
    screenDocFile()
    screenDocxFile()
    # for i in range(0,10):
    #     print(docFiledicList[i])    
    #     # save_doc_to_docx(docFiledicList[0])
    #     word = wc.Dispatch("Word.Application")
    #     doc = word.Documents.Open(docFiledicList[i])
    #     output_file = os.path.splitext(docFiledicList[i])[0] + '.docx'

    #     doc.SaveAs(output_file, 12)  # 12表示docx格式
    #     doc.Close()
    #     word.Quit()
    # print(docFiledicDocxList[0])
    # doc = docx.Document(docFiledicDocxList[0])
    # new_doc = docx.Document()
    # for para in doc.paragraphs:
    #    # 读取段落中的文本内容和样式
    #     text = para.text
    #     style = para.style
    # # 将文本内容添加到新文档中
    #     new_para = new_doc.add_paragraph(text)
    # # 设置新段落的样式
    #     new_para.style = style.name
    new_doc = docx.Document()
    new_doc.add_heading("文档标题",0)
    
    new_doc.save(os.path.splitext(docFiledicList[0])[0]+"_001" + '.docx')
# python inference_main.py -m "logs/44k/G_36800.pth" -c "configs/config.json" -n "you_8.wav" -t 0 -s "lrin"