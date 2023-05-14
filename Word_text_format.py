import docx
import os
# import win32com.client
from docx import Document
import comtypes.client
import comtypes.gen
from win32com import client as wc
import docx2python


docFiledicList=[]
docxFiledicList=[]

def screenDocFile():
# 遍历当前目录下的所有文件夹和子文件夹
    for root, dirs, files in os.walk('.'):
        # 遍历当前文件夹下的所有文件
        for file in files:
            # 检查文件扩展名是否为".doc"或".docx"
            if file.endswith('.doc'):
                # 打开Word文档
                docFiledicList.append(os.path.abspath(os.path.join(root, file)))

                # doc = docx.Document(os.path.join(root, file))
                # 在这里进行其他的操作，例如读取文本、修改文本等

def screenDocxFile():
# 遍历当前目录下的所有文件夹和子文件夹
    for root, dirs, files in os.walk('.'):
        # 遍历当前文件夹下的所有文件
        for file in files:
            # 检查文件扩展名是否为".doc"或".docx"
            if file.endswith('.docx'):
                # 打开Word文档
                docxFiledicList.append(os.path.abspath(os.path.join(root, file)))

                # doc = docx.Document(os.path.join(root, file))
                # 在这里进行其他的操作，例如读取文本、修改文本等


def save_doc_to_docx(rawpath):  # doc转docx
    word = wc.Dispatch("Word.Application")
    doc = word.Documents.Open(docFiledicList[0])
    output_file = os.path.splitext(rawpath)[0] + '.docx'

    doc.SaveAs(output_file, 12)  # 12表示docx格式
    doc.Close()
    word.Quit()
    word.Quit()
    #     output_file = os.path.splitext(docFiledicList[i])[0] + '.docx'

def word():
    print("Hello worlf")

    target_doc = Document('Python-docx.docx')
    source_doc = Document('ex.docx')

    # 将源文档的所有段落、图片和数学公式添加到目标文档中
    for elem in source_doc.element.body:
        # 如果是段落，将其内容添加到目标文档中
        if isinstance(elem, docx.text.paragraph.Paragraph):
            p = docx.text.paragraph.Paragraph(elem, source_doc)
            target_doc.add_paragraph(p.text, style=p.style)
            target_doc.add_paragraph("p.text, style=p.style")
        # # 如果是图片，将其添加到目标文档中
        # elif isinstance(elem, docx.GroupShape):
        #     for item in elem.iter_shape_elms():
        #         if item.get('type') == '#_x0000_t75':
        #             target_doc.add_picture(item.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'])
        # 如果是数学公式，将其添加到目标文档中
        # if isinstance(elem, docx.oxml.math.Paragraph):
        #     p = docx.math.paragraph.MathParagraph(elem, source_doc)
        #     target_doc.add_paragraph(p.text, style=p.style)

# 保存目标文档
    target_doc.save('Python-docx.docx')

    # document.add_heading('文章标题', 0)
    # p = document.add_paragraph('这是一个段落')

    # p.add_run('bold').bold = True           # 添加粗体文字
    # p.add_run(' and some ')                 # 添加默认格式文字
    # p.add_run('italic.').italic = True      # 添加斜体文字
    
    # document.add_heading('这是一级标题, level 1', level=1)

    # document.add_paragraph('这也是一个段落', style='Intense Quote')
    # document.add_picture("1582195139104463.jpg")

    target_doc.save('Python-docx.docx')


if __name__ == '__main__':    
    # screenDocFile()
    # screenDocxFile()    
    
    word()

    # doc = docx.Document()
    # new_doc = docx.Document()    
    
    # new_doc.add_heading("文档标题",0)    
    # new_doc.save(os.path.splitext(docFiledicList[0])[0]+"_001" + '.docx')
# python inference_main.py -m "logs/44k/G_36800.pth" -c "configs/config.json" -n "you_8.wav" -t 0 -s "lrin"